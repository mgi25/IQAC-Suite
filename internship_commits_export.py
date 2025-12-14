import os
import sys
import time
import requests
from datetime import datetime, timedelta
from collections import defaultdict

try:
    from docx import Document
except ImportError:
    print("Missing dependency: python-docx. Install it using:")
    print("    pip install python-docx requests")
    sys.exit(1)

# ---------------- CONFIGURATION ---------------- #

# Edit these values before running
GITHUB_USERNAME = "mgi25"      # your GitHub username (optional)
REPO_OWNER = "mgi25"                     # repo owner (user or org)
REPO_NAME = "IQAC-Suite"                       # repository name
AUTHOR_FILTER = "mgi25"        # username or email used in commits

# Optional date filters (ISO 8601). Use None to disable.
# Example: "2025-04-01T00:00:00Z" or just "2025-04-01"
START_DATE = None
END_DATE = None

OUTPUT_FILENAME = f"internship_commits_{REPO_NAME}.docx"

API_BASE_URL = "https://api.github.com"

# ------------------------------------------------ #

def get_token():
    token = os.getenv("GITHUB_TOKEN")
    if not token:
        print("ERROR: Environment variable GITHUB_TOKEN is not set.")
        print("Create a personal access token on GitHub and set it like:")
        print("  export GITHUB_TOKEN=your_token_here   (Linux/macOS)")
        print("  setx GITHUB_TOKEN your_token_here      (Windows, new terminal needed)")
        sys.exit(1)
    return token

def parse_date(date_str):
    """
    Parse various ISO-like datetime strings to a datetime object.
    """
    if not date_str:
        return None
    # GitHub usually returns e.g. "2025-04-01T12:34:56Z"
    try:
        return datetime.strptime(date_str, "%Y-%m-%dT%H:%M:%SZ")
    except ValueError:
        # Try a simpler date format
        try:
            return datetime.strptime(date_str, "%Y-%m-%d")
        except ValueError:
            return None

def fetch_commits(owner, repo, author, token, since=None, until=None):
    """
    Fetch commits from the GitHub API with pagination and optional filters.
    Also fetch per-commit stats (additions, deletions, files changed).
    """
    headers = {
        "Authorization": f"Bearer {token}",
        "Accept": "application/vnd.github+json"
    }

    all_commits = []
    page = 1

    while True:
        params = {
            "per_page": 100,
            "page": page,
        }
        # author filter (username or email)
        if author:
            params["author"] = author
        # date filters
        if since:
            params["since"] = since
        if until:
            params["until"] = until

        url = f"{API_BASE_URL}/repos/{owner}/{repo}/commits"
        resp = requests.get(url, headers=headers, params=params)

        if resp.status_code == 403:
            # Possibly rate limited
            print("GitHub API returned 403 (forbidden or rate-limited).")
            print("Response:", resp.text)
            break
        elif resp.status_code != 200:
            print(f"Error fetching commits (status {resp.status_code}): {resp.text}")
            break

        commits_page = resp.json()
        if not commits_page:
            # No more commits
            break

        for c in commits_page:
            sha = c.get("sha")
            commit_info = c.get("commit", {})
            author_info = commit_info.get("author", {})
            commit_date_str = author_info.get("date")
            commit_date = parse_date(commit_date_str)
            message = commit_info.get("message", "").strip()
            html_url = c.get("html_url")

            # Fetch detailed stats for this commit
            stats_url = f"{API_BASE_URL}/repos/{owner}/{repo}/commits/{sha}"
            stats_resp = requests.get(stats_url, headers=headers)
            additions = deletions = files_changed = 0

            if stats_resp.status_code == 200:
                stats_data = stats_resp.json()
                stats = stats_data.get("stats", {})
                additions = stats.get("additions", 0)
                deletions = stats.get("deletions", 0)
                files = stats_data.get("files", []) or []
                files_changed = len(files)
            else:
                print(f"Warning: could not fetch stats for commit {sha} (status {stats_resp.status_code})")

            all_commits.append({
                "sha": sha,
                "short_sha": sha[:7] if sha else "",
                "date": commit_date,
                "date_str": commit_date_str,
                "message": message,
                "url": html_url,
                "additions": additions,
                "deletions": deletions,
                "files_changed": files_changed,
            })

            # Be kind to the API (very small delay)
            time.sleep(0.05)

        print(f"Fetched page {page}, {len(commits_page)} commits.")
        page += 1

    # Sort commits chronologically (oldest first)
    all_commits = [c for c in all_commits if c["date"] is not None]
    all_commits.sort(key=lambda x: x["date"])

    return all_commits

def group_commits_by_week(commits):
    """
    Group commits by week (Monday as the start of week).
    Returns a list of (week_start_date, summary_dict).
    summary_dict has: commits, additions, deletions
    """
    weeks = defaultdict(lambda: {"commits": [], "additions": 0, "deletions": 0})

    for c in commits:
        d = c["date"].date()
        # Monday of the week
        monday = d - timedelta(days=d.weekday())
        bucket = weeks[monday]
        bucket["commits"].append(c)
        bucket["additions"] += c["additions"]
        bucket["deletions"] += c["deletions"]

    # Turn into sorted list
    items = sorted(weeks.items(), key=lambda x: x[0])
    return items

def create_docx(commits, weekly_summaries, owner, repo, author, output_filename):
    if not commits:
        print("No commits found. Document will still be created but with minimal content.")

    doc = Document()

    # ---------------- Title Page ----------------
    doc.add_heading(f"Internship Commit Log – {repo}", level=0)
    doc.add_paragraph(f"GitHub Contributions by {author}")
    doc.add_paragraph("")
    now_str = datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")
    doc.add_paragraph(f"Generated on: {now_str}")
    doc.add_paragraph(f"Repository: https://github.com/{owner}/{repo}")

    doc.add_page_break()

    # ---------------- Section 1: Overview ----------------
    doc.add_heading("1. Overview", level=1)

    total_commits = len(commits)
    total_additions = sum(c["additions"] for c in commits)
    total_deletions = sum(c["deletions"] for c in commits)

    if commits:
        start_date = commits[0]["date"].strftime("%Y-%m-%d")
        end_date = commits[-1]["date"].strftime("%Y-%m-%d")
    else:
        start_date = end_date = "N/A"

    p = doc.add_paragraph()
    p.add_run("Repository: ").bold = True
    p.add_run(f"{owner}/{repo}\n")
    p.add_run("Author filter: ").bold = True
    p.add_run(f"{author}\n")
    p.add_run("Total commits: ").bold = True
    p.add_run(str(total_commits) + "\n")
    p.add_run("Date range: ").bold = True
    p.add_run(f"{start_date} to {end_date}\n")
    p.add_run("Total additions: ").bold = True
    p.add_run(str(total_additions) + "\n")
    p.add_run("Total deletions: ").bold = True
    p.add_run(str(total_deletions) + "\n")

    # ---------------- Section 2: Weekly Summary ----------------
    doc.add_heading("2. Weekly Summary", level=1)

    if not weekly_summaries:
        doc.add_paragraph("No weekly data available (no commits).")
    else:
        for week_start, summary in weekly_summaries:
            week_label = week_start.strftime("%Y-%m-%d")
            doc.add_heading(f"Week of {week_label}", level=2)
            commits_count = len(summary["commits"])
            additions = summary["additions"]
            deletions = summary["deletions"]

            bullet = doc.add_paragraph(style="List Bullet")
            bullet.add_run(f"Number of commits: {commits_count}")
            bullet = doc.add_paragraph(style="List Bullet")
            bullet.add_run(f"Total additions: {additions}")
            bullet = doc.add_paragraph(style="List Bullet")
            bullet.add_run(f"Total deletions: {deletions}")

    doc.add_page_break()

    # ---------------- Section 3: Detailed Commit Log ----------------
    doc.add_heading("3. Detailed Commit Log", level=1)

    if not commits:
        doc.add_paragraph("No commits found for the given filters.")
    else:
        for idx, c in enumerate(commits, start=1):
            date_str = c["date"].strftime("%Y-%m-%d %H:%M:%S UTC")
            short_sha = c["short_sha"]
            additions = c["additions"]
            deletions = c["deletions"]
            files_changed = c["files_changed"]
            url = c["url"] or ""

            # Commit message: first line bold, rest normal
            message = c["message"] or ""
            lines = message.splitlines()
            first_line = lines[0] if lines else ""
            body_lines = lines[1:] if len(lines) > 1 else []

            # Commit header paragraph
            header_p = doc.add_paragraph()
            header_p.add_run(f"Commit {idx}: ").bold = True
            header_p.add_run(first_line if first_line else "(no message)")

            # Metadata paragraph
            meta_p = doc.add_paragraph()
            meta_p.add_run("Date: ").bold = True
            meta_p.add_run(f"{date_str}  ")
            meta_p.add_run("SHA: ").bold = True
            meta_p.add_run(f"{short_sha}\n")
            meta_p.add_run("Additions: ").bold = True
            meta_p.add_run(f"{additions}  ")
            meta_p.add_run("Deletions: ").bold = True
            meta_p.add_run(f"{deletions}  ")
            meta_p.add_run("Files changed: ").bold = True
            meta_p.add_run(f"{files_changed}\n")
            if url:
                meta_p.add_run("URL: ").bold = True
                meta_p.add_run(url)

            # Body (if any)
            if body_lines:
                body_text = "\n".join(body_lines)
                doc.add_paragraph(body_text)

            # Spacer between commits
            doc.add_paragraph("")

    # ---------------- Save Document ----------------
    doc.save(output_filename)
    print(f"Export complete. File saved as: {output_filename}")

def main():
    token = get_token()

    since = START_DATE
    until = END_DATE

    commits = fetch_commits(
        owner=REPO_OWNER,
        repo=REPO_NAME,
        author=AUTHOR_FILTER,
        token=token,
        since=since,
        until=until,
    )

    print(f"Total commits fetched: {len(commits)}")

    weekly_summaries = group_commits_by_week(commits)

    create_docx(
        commits=commits,
        weekly_summaries=weekly_summaries,
        owner=REPO_OWNER,
        repo=REPO_NAME,
        author=AUTHOR_FILTER,
        output_filename=OUTPUT_FILENAME,
    )

if __name__ == "__main__":
    main()
